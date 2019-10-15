import pprint
import urllib.parse
import qrcode
import requests
from docx import Document
import string
import pathlib


def findBigWords(t):
    words = t.split()
    keys = []
    logText = ''

    for i in range(len(words)):
        if len(words[i]) < 5:
            logText += words[i] + " is too short./n"
        else:
            logText += words[i] + " is a good key./n"
            if not words[i].lower() in keys:
                keys.append(words[i].lower().translate(str.maketrans('', '', string.punctuation)))
                logText += "Added " + words[i] + " to list of keys./n"
            else:
                logText += words[i] + " is already in the list of keys. SKipping.../n"
    return keys


def getJokes(path):
    jokes = []
    # define the path
    currentDirectory = pathlib.Path(path)

    for currentFile in currentDirectory.iterdir():
        f = open(currentFile, 'r')
        content = f.read()
        jokes.append(content)
    return jokes


def encryptJoke(j, k):
    # defining the api-endpoint
    API_ENDPOINT = "https://webserver.myhhgttg.com/AES/AES_api.php"

    # data to be sent to api
    postObj = {'key': key,
               'data': joke,
               'encBit': '128',
               'operation': 'encrypt'}

    # sending post request and saving response as response object
    r = requests.post(url=API_ENDPOINT, data=postObj)

    # extracting response text
    return r.text


def getQRCode(url, fileName):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=5,
        border=4,
    )

    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(fileName)


def createDOCX(key, file, qr, url,testingKey):
    document = Document()
    document.add_paragraph('Your key is: ' + key)
    document.add_paragraph(
        'Each guest has been provided a key and an encrypted message. The QR code below links to a max.gov website '
        'that will let you decode the message.  Your key cannot decode your message. To find the message you must '
        'mingle with other guests and try their key against your message and vice versa.')
    document.add_paragraph(
        'To try and decode the message, open the camera app in your phone and focus it on the QR code. Your phone '
        'should ask if you want to open the link in Safari.  Once the link opens the encrypted message will be '
        'populated for you. You will need to enter the correct key and click press the decrypt button. If you have '
        'the right key the message will be displayed below the button.')
    document.add_picture(qr)
    document.add_paragraph(url)
    if len(testingKey) > 0:
        document.add_paragraph('The key for this cyphertext is: ' + testingKey)
    document.save(file)


# Text block to get keys from
text = "We the People of the United States, in Order to form a more perfect Union, establish Justice, insure domestic " \
       "Tranquility, provide for the common defence, promote the general Welfare, and secure the Blessings of Liberty " \
       "to ourselves and our Posterity, do ordain and establish this Constitution for the United States of America. "

keysList = findBigWords(text)
jokesList = getJokes('data/jokes/')

for i in range(len(jokesList)):
    key = keysList[i]
    joke = jokesList[i]
    wordFile = 'data/handouts/handout' + str(i) + '.docx'
    qrFile = 'data/qrCodes/qrImage' + str(i) + '.png'
    # the key variable is going to be used to encrypt the joke
    # We want to get a different key to give each guest
    printKey = i + 1
    lastJoke = len(jokesList)
    if printKey <= lastJoke:
        printKey = keysList[printKey]
    else:
        printKey = keysList[0]

    # get the URL with the cypher tex as a parameter so people don't have to type in the long string
    baseURL = 'https://webserver.myhhgttg.com/AES/?q='
    ct = urllib.parse.quote(encryptJoke(joke, key))
    printURL = baseURL + ct
    testingKey = key

    # get a QR code for the printURL
    getQRCode(printURL, qrFile)

    # create the docx handout...
    # createDOCX(printKey, wordFile, qrFile, printURL)

    createDOCX(printKey, wordFile, qrFile, printURL, testingKey)
